"""
Document Views - Views të avancuara për menaxhimin e dokumenteve
Integron template engine, workflow system, signatures dhe AI features
"""

import json
import logging
from typing import Dict, List, Any, Optional
from datetime import datetime, timedelta

from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import LoginRequiredMixin, UserPassesTestMixin
from django.views.generic import (
    CreateView, ListView, DetailView, UpdateView, DeleteView, FormView, TemplateView
)
from django.views.decorators.http import require_http_methods, require_POST
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse, HttpResponse, Http404, HttpResponseForbidden
from django.core.exceptions import ValidationError, PermissionDenied
from django.core.paginator import Paginator
from django.contrib import messages
from django.urls import reverse_lazy, reverse
from django.db import transaction
from django.db.models import Q, Count, Prefetch
from django.utils import timezone
from django.conf import settings

from ..models.document_models import (
    Document, DocumentTemplate, DocumentType, DocumentStatus,
    DocumentComment, DocumentVersion, LLMInteraction, DocumentAuditLog
)
from ..services.document_service import DocumentEditingService
from ..services.llm_service import LegalLLMService, DocumentContext
from ..advanced_features.template_engine import LegalTemplateEngine, TemplateContext
from ..advanced_features.workflow_system import WorkflowEngine
from ..advanced_features.signature_system import SignatureService, SignerInfo
from ..advanced_features.document_automation import DocumentAutomationEngine
from ..forms import (
    DocumentForm, DocumentCommentForm, DocumentUploadForm, 
    TemplateVariableForm, SignatureRequestForm
)

logger = logging.getLogger(__name__)

class DocumentListView(LoginRequiredMixin, ListView):
    """Lista e dokumenteve me filtrime të avancuara"""
    model = Document
    template_name = 'document_editor/documents/list.html'
    context_object_name = 'documents'
    paginate_by = 20

    def get_queryset(self):
        queryset = Document.objects.select_related(
            'case', 'document_type', 'status', 'owned_by', 'created_by'
        ).prefetch_related(
            'editors', 'comments', 'signatures'
        )

        # Filter by user permissions
        user = self.request.user
        if user.role == 'client':
            queryset = queryset.filter(case__client__user=user)
        elif user.role in ['lawyer', 'paralegal']:
            queryset = queryset.filter(
                Q(owned_by=user) | 
                Q(assigned_to=user) | 
                Q(editors__user=user) |
                Q(case__assigned_to=user)
            ).distinct()

        # Apply search filters
        search_query = self.request.GET.get('search')
        if search_query:
            queryset = queryset.filter(
                Q(title__icontains=search_query) |
                Q(content__icontains=search_query) |
                Q(case__title__icontains=search_query)
            )

        # Apply filters
        document_type = self.request.GET.get('type')
        if document_type:
            queryset = queryset.filter(document_type_id=document_type)

        status = self.request.GET.get('status')
        if status:
            queryset = queryset.filter(status_id=status)

        case_type = self.request.GET.get('case_type')
        if case_type:
            queryset = queryset.filter(case__case_type=case_type)

        # Sort
        sort_by = self.request.GET.get('sort', '-updated_at')
        queryset = queryset.order_by(sort_by)

        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        
        # Add filter options
        context['document_types'] = DocumentType.objects.all()
        context['document_statuses'] = DocumentStatus.objects.all()
        context['case_types'] = [
            ('civil', 'Civil'),
            ('criminal', 'Criminal'), 
            ('commercial', 'Commercial'),
            ('family', 'Family')
        ]
        
        # Add current filters
        context['current_filters'] = {
            'search': self.request.GET.get('search', ''),
            'type': self.request.GET.get('type', ''),
            'status': self.request.GET.get('status', ''),
            'case_type': self.request.GET.get('case_type', ''),
            'sort': self.request.GET.get('sort', '-updated_at')
        }

        # Add statistics
        context['stats'] = {
            'total_documents': context['documents'].count() if hasattr(context['documents'], 'count') else len(context['documents']),
            'draft_documents': Document.objects.filter(status__name='Draft').count(),
            'pending_review': Document.objects.filter(status__name='Review').count(),
            'signed_documents': Document.objects.filter(signatures__isnull=False).distinct().count()
        }

        return context

class DocumentDetailView(LoginRequiredMixin, DetailView):
    """Detajet e dokumentit me të gjitha features"""
    model = Document
    template_name = 'document_editor/documents/detail.html'
    context_object_name = 'document'

    def get_queryset(self):
        return Document.objects.select_related(
            'case', 'document_type', 'status', 'owned_by', 'created_by',
            'locked_by', 'last_edited_by'
        ).prefetch_related(
            'editors__user',
            'comments__author',
            'signatures__signer',
            'version_history__created_by',
            'llm_interactions__user',
            'audit_logs__user'
        )

    def get_object(self, queryset=None):
        obj = super().get_object(queryset)
        
        # Check permissions
        if not self._can_view_document(obj):
            raise Http404("Document not found")

        # Log view action
        DocumentAuditLog.objects.create(
            document=obj,
            user=self.request.user,
            action='view',
            ip_address=self._get_client_ip(),
            user_agent=self.request.META.get('HTTP_USER_AGENT', '')
        )

        return obj

    def _can_view_document(self, document):
        """Check if user can view document"""
        user = self.request.user
        
        if user.role == 'admin':
            return True
        elif user.role == 'client':
            return document.case.client.user == user
        elif user.role in ['lawyer', 'paralegal']:
            return (
                document.owned_by == user or
                document.created_by == user or
                document.editors.filter(user=user).exists() or
                document.case.assigned_to == user
            )
        
        return False

    def _get_client_ip(self):
        """Get client IP address"""
        x_forwarded_for = self.request.META.get('HTTP_X_FORWARDED_FOR')
        if x_forwarded_for:
            ip = x_forwarded_for.split(',')[0]
        else:
            ip = self.request.META.get('REMOTE_ADDR')
        return ip

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        document = self.object

        # Add editing capabilities
        context['can_edit'] = document.can_edit(self.request.user)
        context['is_locked'] = document.is_locked
        context['locked_by_current_user'] = document.locked_by == self.request.user

        # Add comments
        context['comments'] = document.comments.filter(
            parent_comment__isnull=True
        ).select_related('author').prefetch_related('replies__author')
        context['comment_form'] = DocumentCommentForm()

        # Add version history
        context['versions'] = document.version_history.select_related('created_by')[:10]

        # Add workflow info
        if hasattr(document, 'workflow'):
            context['workflow'] = document.workflow
            context['workflow_steps'] = document.workflow.steps.select_related(
                'assigned_users'
            ).prefetch_related('actions__user')

        # Add signature info
        context['signatures'] = document.signatures.select_related('signer')
        context['signature_requests'] = document.signature_requests.all()

        # Add LLM interaction history
        context['llm_interactions'] = document.llm_interactions.select_related('user')[:5]

        # Add related documents
        context['related_documents'] = Document.objects.filter(
            case=document.case
        ).exclude(id=document.id)[:5]

        # Add document stats
        context['document_stats'] = {
            'word_count': len(document.content.split()) if document.content else 0,
            'character_count': len(document.content) if document.content else 0,
            'version_count': document.version_history.count(),
            'comment_count': document.comments.count(),
            'edit_count': document.metadata.get('edit_count', 0) if document.metadata else 0
        }

        return context

class DocumentCreateView(LoginRequiredMixin, UserPassesTestMixin, CreateView):
    """Krijimi i dokumentit të ri me AI assistance"""
    model = Document
    form_class = DocumentForm
    template_name = 'document_editor/documents/create.html'

    def test_func(self):
        return self.request.user.role in ['admin', 'lawyer', 'paralegal']

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs['user'] = self.request.user
        return kwargs

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        
        # Add template suggestions if case is provided
        case_id = self.request.GET.get('case_id')
        if case_id:
            try:
                from cases.models import Case
                case = Case.objects.get(id=case_id)
                context['case'] = case
                
                # Get AI template suggestions
                automation_engine = DocumentAutomationEngine()
                suggestions = automation_engine.suggest_document_template(
                    case_info={
                        'title': case.title,
                        'description': case.description,
                        'case_type': case.case_type
                    }
                )
                context['template_suggestions'] = suggestions[:5]
                
            except Case.DoesNotExist:
                pass

        # Add available templates
        context['templates'] = DocumentTemplate.objects.filter(is_active=True)
        
        # Add document types
        context['document_types'] = DocumentType.objects.all()
        
        return context

    def form_valid(self, form):
        document = form.save(commit=False)
        document.created_by = self.request.user
        document.owned_by = self.request.user
        
        # Generate content from template if selected
        template_id = self.request.POST.get('template_id')
        if template_id:
            try:
                template = DocumentTemplate.objects.get(id=template_id)
                variables = json.loads(self.request.POST.get('template_variables', '{}'))
                
                template_engine = LegalTemplateEngine()
                context = TemplateContext(variables=variables)
                document.content = template_engine.render_template(template, context)
                document.template_used = template
                
            except (DocumentTemplate.DoesNotExist, json.JSONDecodeError, ValidationError) as e:
                messages.error(self.request, f"Template error: {str(e)}")

        document.save()
        
        # Create workflow if applicable
        try:
            workflow_engine = WorkflowEngine()
            workflow_engine.create_workflow(document)
        except ValidationError:
            pass  # No workflow template found
        
        messages.success(self.request, f"Document '{document.title}' created successfully!")
        return redirect('document_editor:document_detail', pk=document.pk)

class DocumentUpdateView(LoginRequiredMixin, UserPassesTestMixin, UpdateView):
    """Editimi i dokumentit"""
    model = Document
    form_class = DocumentForm
    template_name = 'document_editor/documents/edit.html'

    def test_func(self):
        document = self.get_object()
        return document.can_edit(self.request.user)

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs['user'] = self.request.user
        return kwargs

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        document = self.object

        # Add editing service
        context['editing_service'] = DocumentEditingService()
        
        # Add AI suggestions
        context['ai_enabled'] = getattr(settings, 'LLM_ENABLED', False)
        
        # Add collaboration info
        context['active_editors'] = document.editors.filter(
            user__last_login__gte=timezone.now() - timedelta(hours=1)
        )
        
        return context

    def form_valid(self, form):
        document = form.save(commit=False)
        
        # Use document editing service for saving
        editing_service = DocumentEditingService()
        
        try:
            editing_service.save_document_content(
                document=document,
                content=form.cleaned_data['content'],
                content_html=form.cleaned_data.get('content_html', ''),
                user=self.request.user,
                create_version='create_version' in self.request.POST
            )
            
            messages.success(self.request, "Document updated successfully!")
            
        except (ValidationError, PermissionDenied) as e:
            messages.error(self.request, str(e))
            return self.form_invalid(form)

        return redirect('document_editor:document_detail', pk=document.pk)

class DocumentDeleteView(LoginRequiredMixin, UserPassesTestMixin, DeleteView):
    """Fshirja e dokumentit"""
    model = Document
    template_name = 'document_editor/documents/confirm_delete.html'
    success_url = reverse_lazy('document_editor:document_list')

    def test_func(self):
        document = self.get_object()
        return (
            self.request.user.role == 'admin' or 
            document.owned_by == self.request.user
        )

    def delete(self, request, *args, **kwargs):
        document = self.get_object()
        
        # Log deletion
        DocumentAuditLog.objects.create(
            document=document,
            user=request.user,
            action='delete',
            details=f"Document '{document.title}' deleted"
        )
        
        messages.success(request, f"Document '{document.title}' deleted successfully!")
        return super().delete(request, *args, **kwargs)

@login_required
@require_POST
def document_lock_toggle(request, document_id):
    """Toggle document lock status"""
    try:
        document = get_object_or_404(Document, id=document_id)
        
        if not document.can_edit(request.user):
            return JsonResponse({'success': False, 'error': 'Permission denied'})
        
        if document.is_locked:
            success = document.unlock_document(request.user)
            action = 'unlocked'
        else:
            success = document.lock_document(request.user)
            action = 'locked'
        
        if success:
            return JsonResponse({
                'success': True, 
                'action': action,
                'is_locked': document.is_locked,
                'locked_by': document.locked_by.username if document.locked_by else None
            })
        else:
            return JsonResponse({'success': False, 'error': f'Cannot {action} document'})
            
    except Exception as e:
        logger.error(f"Document lock toggle error: {e}")
        return JsonResponse({'success': False, 'error': str(e)})

@login_required
@require_POST
def document_comment_add(request, document_id):
    """Add comment to document"""
    try:
        document = get_object_or_404(Document, id=document_id)
        form = DocumentCommentForm(request.POST)
        
        if form.is_valid():
            editing_service = DocumentEditingService()
            comment = editing_service.add_comment(
                document=document,
                content=form.cleaned_data['content'],
                user=request.user,
                position_start=form.cleaned_data.get('position_start'),
                position_end=form.cleaned_data.get('position_end'),
                selected_text=form.cleaned_data.get('selected_text', ''),
                parent_comment_id=form.cleaned_data.get('parent_comment_id')
            )
            
            return JsonResponse({
                'success': True,
                'comment_id': comment.id,
                'comment_html': render(request, 'document_editor/partials/comment.html', {
                    'comment': comment
                }).content.decode()
            })
        else:
            return JsonResponse({'success': False, 'errors': form.errors})
            
    except Exception as e:
        logger.error(f"Add comment error: {e}")
        return JsonResponse({'success': False, 'error': str(e)})

@login_required
@require_POST
def document_comment_resolve(request, comment_id):
    """Resolve a comment"""
    try:
        editing_service = DocumentEditingService()
        comment = editing_service.resolve_comment(comment_id, request.user)
        
        return JsonResponse({
            'success': True,
            'comment_id': comment.id,
            'resolved_at': comment.resolved_at.isoformat(),
            'resolved_by': comment.resolved_by.username
        })
        
    except ValidationError as e:
        return JsonResponse({'success': False, 'error': str(e)})
    except Exception as e:
        logger.error(f"Resolve comment error: {e}")
        return JsonResponse({'success': False, 'error': str(e)})

@login_required
@require_POST
def document_version_restore(request, document_id, version_number):
    """Restore a document version"""
    try:
        document = get_object_or_404(Document, id=document_id)
        
        if not document.can_edit(request.user):
            return JsonResponse({'success': False, 'error': 'Permission denied'})
        
        editing_service = DocumentEditingService()
        editing_service.restore_document_version(document, version_number, request.user)
        
        messages.success(request, f"Document restored to version {version_number}")
        
        return JsonResponse({
            'success': True,
            'redirect_url': reverse('document_editor:document_detail', kwargs={'pk': document.pk})
        })
        
    except (ValidationError, PermissionDenied) as e:
        return JsonResponse({'success': False, 'error': str(e)})
    except Exception as e:
        logger.error(f"Version restore error: {e}")
        return JsonResponse({'success': False, 'error': str(e)})

@login_required
def document_ai_suggestions(request, document_id):
    """Get AI suggestions for document improvement"""
    try:
        document = get_object_or_404(Document, id=document_id)
        
        if not document.can_edit(request.user):
            return JsonResponse({'success': False, 'error': 'Permission denied'})
        
        llm_service = LegalLLMService()
        context = DocumentContext(
            title=document.title,
            content=document.content,
            document_type=document.document_type.name,
            case_type=getattr(document.case, 'case_type', None)
        )
        
        suggestion_type = request.GET.get('type', 'improvements')
        
        if suggestion_type == 'improvements':
            response = llm_service.suggest_improvements(context)
        elif suggestion_type == 'review':
            response = llm_service.review_document(context)
        elif suggestion_type == 'compliance':
            response = llm_service.analyze_legal_compliance(context)
        else:
            return JsonResponse({'success': False, 'error': 'Invalid suggestion type'})
        
        # Save interaction
        LLMInteraction.objects.create(
            document=document,
            user=request.user,
            interaction_type=suggestion_type,
            prompt=f"Get {suggestion_type} for document",
            llm_response=response.text if not response.error else response.error,
            confidence_score=response.confidence,
            processing_time=response.processing_time,
            llm_model=response.model_used,
            llm_provider=response.provider
        )
        
        if response.error:
            return JsonResponse({'success': False, 'error': response.error})
        
        return JsonResponse({
            'success': True,
            'suggestions': response.text,
            'confidence': response.confidence,
            'processing_time': response.processing_time
        })
        
    except Exception as e:
        logger.error(f"AI suggestions error: {e}")
        return JsonResponse({'success': False, 'error': str(e)})

@login_required
@require_POST
def document_ai_generate_content(request, document_id):
    """Generate content using AI"""
    try:
        document = get_object_or_404(Document, id=document_id)
        
        if not document.can_edit(request.user):
            return JsonResponse({'success': False, 'error': 'Permission denied'})
        
        automation_engine = DocumentAutomationEngine()
        
        # Get generation parameters
        content_type = request.POST.get('content_type', 'expansion')
        instructions = request.POST.get('instructions', '')
        
        case_info = {
            'title': document.case.title,
            'description': document.case.description,
            'case_type': getattr(document.case, 'case_type', 'civil')
        }
        
        # Generate content based on type
        if content_type == 'template':
            template_id = request.POST.get('template_id')
            variables = json.loads(request.POST.get('variables', '{}'))
            
            result = automation_engine.generate_document_content(
                template_id=int(template_id),
                variables=variables,
                case_info=case_info,
                enhancement_level=request.POST.get('enhancement_level', 'standard')
            )
        else:
            # Custom generation based on instructions
            llm_service = LegalLLMService()
            context = DocumentContext(
                title=document.title,
                content=document.content,
                document_type=document.document_type.name,
                case_type=case_info['case_type']
            )
            
            response = llm_service.suggest_improvements(context)
            result = {
                'success': not response.error,
                'content': response.text if not response.error else response.error
            }
        
        if result['success']:
            # Save interaction
            LLMInteraction.objects.create(
                document=document,
                user=request.user,
                interaction_type='generate',
                prompt=f"Generate {content_type}: {instructions}",
                llm_response=result['content'],
                llm_model='automation_engine'
            )
            
            return JsonResponse({
                'success': True,
                'content': result['content'],
                'template_used': result.get('template_used'),
                'enhancement_level': result.get('enhancement_level')
            })
        else:
            return JsonResponse({'success': False, 'error': result.get('error', 'Generation failed')})
        
    except Exception as e:
        logger.error(f"AI content generation error: {e}")
        return JsonResponse({'success': False, 'error': str(e)})

@login_required
def document_export(request, document_id):
    """Export document in various formats"""
    try:
        document = get_object_or_404(Document, id=document_id)
        
        if not document.can_edit(request.user):
            raise PermissionDenied("You don't have permission to export this document")
        
        export_format = request.GET.get('format', 'pdf')
        
        # Log export action
        DocumentAuditLog.objects.create(
            document=document,
            user=request.user,
            action='export',
            details=f"Document exported as {export_format}",
            metadata={'format': export_format}
        )
        
        if export_format == 'pdf':
            return _export_as_pdf(document)
        elif export_format == 'docx':
            return _export_as_docx(document)
        elif export_format == 'html':
            return _export_as_html(document)
        else:
            return JsonResponse({'success': False, 'error': 'Unsupported format'})
            
    except Exception as e:
        logger.error(f"Document export error: {e}")
        return JsonResponse({'success': False, 'error': str(e)})

def _export_as_pdf(document):
    """Export document as PDF"""
    try:
        import weasyprint
        from django.template.loader import render_to_string
        
        html_content = render_to_string('document_editor/exports/pdf_template.html', {
            'document': document
        })
        
        pdf = weasyprint.HTML(string=html_content).write_pdf()
        
        response = HttpResponse(pdf, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{document.title}.pdf"'
        return response
        
    except ImportError:
        return JsonResponse({'success': False, 'error': 'PDF export not available'})

def _export_as_docx(document):
    """Export document as DOCX"""
    try:
        from docx import Document as DocxDocument
        from django.http import HttpResponse
        import io
        
        doc = DocxDocument()
        doc.add_heading(document.title, 0)
        
        # Add document content
        paragraphs = document.content.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        # Save to memory
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        response = HttpResponse(
            file_stream.read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{document.title}.docx"'
        return response
        
    except ImportError:
        return JsonResponse({'success': False, 'error': 'DOCX export not available'})

def _export_as_html(document):
    """Export document as HTML"""
    from django.template.loader import render_to_string
    
    html_content = render_to_string('document_editor/exports/html_template.html', {
        'document': document
    })
    
    response = HttpResponse(html_content, content_type='text/html')
    response['Content-Disposition'] = f'attachment; filename="{document.title}.html"'
    return response
